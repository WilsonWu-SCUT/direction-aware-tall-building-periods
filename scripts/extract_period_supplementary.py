#!/usr/bin/env python3
"""Extract the tall-building period tables from the Word supplement."""

from __future__ import annotations

import argparse
import csv
from decimal import Decimal, InvalidOperation
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUTPUT = ROOT / "data" / "period_prediction" / "tall_building_periods.csv"

CSV_FIELDS = [
    "model_id",
    "structural_system_code",
    "structural_system",
    "source_table",
    "source_row",
    "building_height_m",
    "transverse_effective_width_m",
    "longitudinal_effective_width_m",
    "building_function",
    "seismic_intensity_degree",
    "transverse_period_s",
    "longitudinal_period_s",
]

EXPECTED_HEADERS = [
    "ID",
    "Building Height (m)",
    "Transverse Effective Width (m)",
    "Longitudinal Effective Width (m)",
    "Function",
    "Seismic Intensity",
    "Transverse Translational Period (s)",
    "Longitudinal Translational Period (s)",
]

TABLE_CONFIG = [
    ("SW", "Shear wall", "Table S1", 818),
    ("FSW", "Frame-shear wall", "Table S2", 181),
    ("FT", "Frame-tube", "Table S3", 334),
]


def compact_text(value: str) -> str:
    """Collapse Word line breaks and repeated whitespace."""
    return " ".join(value.split())


def require_decimal(value: str, field_name: str, model_id: str) -> str:
    """Validate a numeric string while preserving its displayed precision."""
    try:
        parsed = Decimal(value)
    except InvalidOperation as error:
        raise ValueError(
            f"Invalid {field_name} value for {model_id}: {value!r}"
        ) from error
    if parsed <= 0:
        raise ValueError(f"{field_name} must be positive for {model_id}.")
    return value


def parse_intensity(value: str, model_id: str) -> int:
    """Convert source labels such as '7-Degree' to an integer category."""
    prefix = value.split("-", maxsplit=1)[0]
    try:
        intensity = int(prefix)
    except ValueError as error:
        raise ValueError(
            f"Invalid seismic intensity for {model_id}: {value!r}"
        ) from error
    if intensity not in {6, 7, 8}:
        raise ValueError(f"Unsupported seismic intensity for {model_id}: {intensity}")
    return intensity


def extract_records(document_path: Path) -> list[dict[str, object]]:
    """Return validated records from the three supplementary tables."""
    document = Document(document_path)
    if len(document.tables) != len(TABLE_CONFIG):
        raise ValueError(
            f"Expected {len(TABLE_CONFIG)} tables, found {len(document.tables)}."
        )

    records: list[dict[str, object]] = []
    seen_ids: set[str] = set()

    for table, (code, system_name, source_table, expected_count) in zip(
        document.tables, TABLE_CONFIG, strict=True
    ):
        headers = [compact_text(cell.text) for cell in table.rows[0].cells]
        if headers != EXPECTED_HEADERS:
            raise ValueError(
                f"Unexpected headers in {source_table}: {headers!r}"
            )

        data_rows = [
            [compact_text(cell.text) for cell in row.cells]
            for row in table.rows[1:]
            if any(compact_text(cell.text) for cell in row.cells)
        ]
        if len(data_rows) != expected_count:
            raise ValueError(
                f"Expected {expected_count} records in {source_table}, "
                f"found {len(data_rows)}."
            )

        for source_row, values in enumerate(data_rows, start=1):
            if len(values) != len(EXPECTED_HEADERS):
                raise ValueError(
                    f"Expected eight values in {source_table} row {source_row}."
                )

            (
                model_id,
                height,
                transverse_width,
                longitudinal_width,
                building_function,
                intensity_label,
                transverse_period,
                longitudinal_period,
            ) = values

            if model_id in seen_ids:
                raise ValueError(f"Duplicate model identifier: {model_id}")
            if not model_id.startswith(f"{code}-"):
                raise ValueError(
                    f"Model identifier {model_id!r} does not match system {code}."
                )
            if building_function not in {"Residential", "Office", "Hotel"}:
                raise ValueError(
                    f"Unsupported building function for {model_id}: "
                    f"{building_function!r}"
                )

            seen_ids.add(model_id)
            records.append(
                {
                    "model_id": model_id,
                    "structural_system_code": code,
                    "structural_system": system_name,
                    "source_table": source_table,
                    "source_row": source_row,
                    "building_height_m": require_decimal(
                        height, "building_height_m", model_id
                    ),
                    "transverse_effective_width_m": require_decimal(
                        transverse_width,
                        "transverse_effective_width_m",
                        model_id,
                    ),
                    "longitudinal_effective_width_m": require_decimal(
                        longitudinal_width,
                        "longitudinal_effective_width_m",
                        model_id,
                    ),
                    "building_function": building_function,
                    "seismic_intensity_degree": parse_intensity(
                        intensity_label, model_id
                    ),
                    "transverse_period_s": require_decimal(
                        transverse_period, "transverse_period_s", model_id
                    ),
                    "longitudinal_period_s": require_decimal(
                        longitudinal_period, "longitudinal_period_s", model_id
                    ),
                }
            )

    if len(records) != 1333:
        raise ValueError(f"Expected 1,333 total records, found {len(records)}.")
    return records


def write_csv(records: list[dict[str, object]], output_path: Path) -> None:
    """Write the canonical UTF-8 CSV with Unix line endings."""
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with output_path.open("w", encoding="utf-8", newline="") as stream:
        writer = csv.DictWriter(stream, fieldnames=CSV_FIELDS, lineterminator="\n")
        writer.writeheader()
        writer.writerows(records)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("document", type=Path, help="Path to Supplementary.docx")
    parser.add_argument(
        "--output",
        type=Path,
        default=DEFAULT_OUTPUT,
        help=f"Output CSV path (default: {DEFAULT_OUTPUT})",
    )
    args = parser.parse_args()

    records = extract_records(args.document)
    write_csv(records, args.output)
    print(f"Wrote {len(records):,} records to {args.output}")


if __name__ == "__main__":
    main()
